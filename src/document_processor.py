"""
Enhanced document processor for Universal RAG application.
This module handles extraction of text from various file formats including
PDF, DOCX, Excel, and images with specialized methods for each type.
"""

import os
import logging
from typing import List, Dict, Any, Optional
import PyPDF2
import docx
import pandas as pd
from PIL import Image
import requests
import zipfile
import shutil

# Try to import pytesseract, but handle gracefully if not available
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
    
    # Try to check if tesseract is in PATH
    try:
        pytesseract.get_tesseract_version()
    except (pytesseract.TesseractNotFoundError, FileNotFoundError):
        # Common Windows installation paths
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\{}\AppData\Local\Tesseract-OCR\tesseract.exe".format(os.getenv('USERNAME', '')),
            r"C:\Tools\Tesseract-OCR\tesseract.exe",
            r".\tesseract\tesseract.exe",
        ]
        
        tesseract_found = False
        for path in possible_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                tesseract_found = True
                break
        
        if not tesseract_found:
            TESSERACT_AVAILABLE = False
            
except ImportError:
    TESSERACT_AVAILABLE = False
    pytesseract = None

from io import BytesIO

class DocumentProcessor:
    """
    Enhanced document processor that handles various file types with specialized methods.
    
    Supported formats:
    - PDF: Extracts text from PDF files
    - DOCX: Extracts text from Microsoft Word documents
    - XLSX: Extracts data from Excel spreadsheets
    - Images (JPG, PNG): Uses OCR to extract text from images
    """
    
    def __init__(self, auto_setup_tesseract: bool = True):
        """Initialize the document processor."""
        self.logger = logging.getLogger(__name__)
        self.supported_extensions = {
            '.pdf': self._process_pdf_file,
            '.docx': self._process_docx_file,
            '.xlsx': self._process_xlsx_file,
            '.jpg': self._process_image_file,
            '.jpeg': self._process_image_file,
            '.png': self._process_image_file
        }
        
        # Try to setup Tesseract if not available
        if auto_setup_tesseract and not TESSERACT_AVAILABLE:
            self._setup_tesseract()
    
    def _setup_tesseract(self):
        """Attempt to download and setup Tesseract OCR."""
        try:
            tesseract_dir = os.path.join(os.getcwd(), "tesseract")
            tesseract_exe = os.path.join(tesseract_dir, "tesseract.exe")
            
            if os.path.exists(tesseract_exe):
                # Tesseract already exists locally
                global TESSERACT_AVAILABLE, pytesseract
                if pytesseract:
                    pytesseract.pytesseract.tesseract_cmd = tesseract_exe
                    TESSERACT_AVAILABLE = True
                    self.logger.info("Using local Tesseract installation")
                return
            
            self.logger.info("Attempting to download Tesseract OCR...")
            
            # Create tesseract directory
            os.makedirs(tesseract_dir, exist_ok=True)
            
            # Download portable Tesseract (this is a simplified approach)
            # Note: In a real scenario, you'd want to download from a reliable source
            tesseract_url = "https://github.com/UB-Mannheim/tesseract/releases/download/v5.3.0.20221214/tesseract-ocr-w64-setup-5.3.0.20221214.exe"
            
            # For now, we'll create a placeholder and inform the user
            self.logger.warning("Tesseract auto-setup not implemented. Please install manually.")
            
        except Exception as e:
            self.logger.warning(f"Could not setup Tesseract automatically: {e}")
    
    def vectorize_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Specialized method to vectorize PDF files.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            Dict[str, Any]: Processed document ready for vectorization
        """
        self.logger.info(f"Vectorizing PDF: {file_path}")
        return self.process_file(file_path)
    
    def vectorize_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Specialized method to vectorize DOCX files.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            Dict[str, Any]: Processed document ready for vectorization
        """
        self.logger.info(f"Vectorizing DOCX: {file_path}")
        return self.process_file(file_path)
    
    def vectorize_xlsx(self, file_path: str) -> Dict[str, Any]:
        """
        Specialized method to vectorize XLSX files.
        
        Args:
            file_path (str): Path to the XLSX file
            
        Returns:
            Dict[str, Any]: Processed document ready for vectorization
        """
        self.logger.info(f"Vectorizing XLSX: {file_path}")
        return self.process_file(file_path)
    
    def vectorize_image(self, file_path: str) -> Dict[str, Any]:
        """
        Specialized method to vectorize image files using OCR.
        
        Args:
            file_path (str): Path to the image file
            
        Returns:
            Dict[str, Any]: Processed document ready for vectorization
        """
        self.logger.info(f"Vectorizing Image: {file_path}")
        return self.process_file(file_path)
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process a single file and extract its content.
        
        Args:
            file_path (str): Path to the file to process
        
        Returns:
            Dict[str, Any]: Dictionary containing extracted content and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            # Get file metadata
            file_size = os.path.getsize(file_path)
            file_name = os.path.basename(file_path)
            
            # Process the file based on its type
            processor_func = self.supported_extensions[file_extension]
            content = processor_func(file_path)
            
            result = {
                'file_name': file_name,
                'file_path': file_path,
                'file_type': file_extension,
                'file_size': file_size,
                'content': content,
                'status': 'success'
            }
            
            self.logger.info(f"Successfully processed: {file_name}")
            return result
            
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            return {
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'file_type': file_extension,
                'content': '',
                'status': 'error',
                'error': str(e)
            }
    
    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path (str): Path to the directory to process
        
        Returns:
            List[Dict[str, Any]]: List of processed file results
        """
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        results = []
        
        for file_name in os.listdir(directory_path):
            file_path = os.path.join(directory_path, file_name)
            
            # Skip directories and unsupported files
            if os.path.isfile(file_path):
                file_extension = os.path.splitext(file_name)[1].lower()
                if file_extension in self.supported_extensions:
                    result = self.process_file(file_path)
                    results.append(result)
        
        return results
    
    def _process_pdf_file(self, file_path: str) -> str:
        """
        Enhanced PDF text extraction with better error handling.
        
        Args:
            file_path (str): Path to the PDF file
        
        Returns:
            str: Extracted text content
        """
        text_content = ""
        
        try:
            self.logger.info(f"Processing PDF: {os.path.basename(file_path)}")
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                self.logger.info(f"PDF has {total_pages} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content += f"\\n=== Page {page_num + 1} of {total_pages} ===\\n"
                            text_content += page_text.strip() + "\\n"
                        else:
                            self.logger.warning(f"No text found on page {page_num + 1}")
                    except Exception as e:
                        self.logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue
                        
            if not text_content.strip():
                return f"PDF file: {os.path.basename(file_path)} (No extractable text found - file may be image-based)"
            
            self.logger.info(f"Successfully extracted {len(text_content)} characters from PDF")
            return text_content.strip()
                        
        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {e}")
            raise Exception(f"Error reading PDF: {e}")
    
    def _process_docx_file(self, file_path: str) -> str:
        """
        Enhanced DOCX text extraction with tables and formatting.
        
        Args:
            file_path (str): Path to the DOCX file
        
        Returns:
            str: Extracted text content
        """
        try:
            self.logger.info(f"Processing DOCX: {os.path.basename(file_path)}")
            
            doc = docx.Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content += paragraph.text.strip() + "\\n"
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                text_content += f"\\n=== Table {table_count + 1} ===\\n"
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_content += f"Row {row_idx + 1}: {row_text}\\n"
                table_count += 1
            
            if not text_content.strip():
                return f"DOCX file: {os.path.basename(file_path)} (No text content found)"
            
            self.logger.info(f"Successfully extracted {len(text_content)} characters from DOCX ({paragraph_count} paragraphs, {table_count} tables)")
            return text_content.strip()
            
        except Exception as e:
            self.logger.error(f"Error reading DOCX {file_path}: {e}")
            raise Exception(f"Error reading DOCX: {e}")
    
    def _process_xlsx_file(self, file_path: str) -> str:
        """
        Enhanced Excel data extraction with row-based semantic chunking for optimal RAG performance.
        Each row becomes a searchable semantic unit with full context.
        
        Args:
            file_path (str): Path to the Excel file
        
        Returns:
            str: Extracted content optimized for RAG with row-level semantic units
        """
        try:
            self.logger.info(f"Processing XLSX: {os.path.basename(file_path)}")
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_content = ""
            total_rows = 0
            
            self.logger.info(f"Excel file has {len(excel_file.sheet_names)} sheets")
            
            for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Skip empty sheets
                    if df.empty:
                        self.logger.info(f"Sheet '{sheet_name}' is empty, skipping")
                        continue
                    
                    # Clean column names and data
                    df.columns = [str(col).strip() for col in df.columns]
                    df = df.dropna(how='all').reset_index(drop=True)
                    
                    if df.empty:
                        continue
                    
                    text_content += f"\\n=== SHEET: {sheet_name} ===\\n"
                    text_content += f"Data Source: {sheet_name} sheet from {os.path.basename(file_path)}\\n"
                    text_content += f"Contains {len(df)} records with {len(df.columns)} fields\\n"
                    text_content += f"Fields: {', '.join(df.columns)}\\n\\n"
                    
                    # Process each row as a complete semantic unit
                    for idx, row in df.iterrows():
                        # Create a comprehensive row description
                        row_text = f"RECORD {idx + 1} from {sheet_name}: "
                        
                        # Build a natural language description of the row
                        row_attributes = []
                        for col, val in row.items():
                            if pd.notna(val) and str(val).strip():
                                clean_val = str(val).strip()
                                # Create natural language attributes
                                if col.lower() in ['name', 'city', 'state', 'district', 'location']:
                                    row_attributes.append(f"located in {clean_val}")
                                elif col.lower() in ['population', 'pop', 'total_population']:
                                    row_attributes.append(f"has population of {clean_val}")
                                elif col.lower() in ['literacy', 'literacy_rate']:
                                    row_attributes.append(f"literacy rate is {clean_val}%")
                                elif col.lower() in ['rank', 'ranking']:
                                    row_attributes.append(f"ranked {clean_val}")
                                elif col.lower() in ['area', 'land_area']:
                                    row_attributes.append(f"covers area of {clean_val}")
                                elif col.lower() in ['density', 'population_density']:
                                    row_attributes.append(f"population density of {clean_val}")
                                else:
                                    # Generic attribute
                                    row_attributes.append(f"{col.lower().replace('_', ' ')} is {clean_val}")
                        
                        if row_attributes:
                            # Create a complete sentence for each row
                            row_text += ". ".join(row_attributes) + "."
                            
                            # Add contextual information for better searchability
                            row_text += f" [Sheet: {sheet_name}, Row: {idx + 1}]"
                            
                            # Add searchable keywords
                            keywords = []
                            for col, val in row.items():
                                if pd.notna(val) and str(val).strip():
                                    keywords.append(str(val).strip())
                            
                            if keywords:
                                row_text += f" Keywords: {', '.join(keywords[:10])}"  # Limit keywords
                            
                            text_content += row_text + "\\n\\n"
                    
                    total_rows += len(df)
                    
                    # Add sheet-level analytics for context
                    text_content += f"\\n--- {sheet_name} ANALYTICS ---\\n"
                    
                    # Numeric summaries
                    numeric_columns = df.select_dtypes(include=['number']).columns
                    for col in numeric_columns:
                        series = df[col].dropna()
                        if not series.empty and len(series) > 0:
                            text_content += f"STATISTIC: {col} in {sheet_name} ranges from {series.min()} to {series.max()}, average {series.mean():.2f}. "
                            
                            # Add top/bottom records for context
                            if col.lower() in ['population', 'pop']:
                                top_record = df.loc[series.idxmax()]
                                bottom_record = df.loc[series.idxmin()]
                                name_col = next((c for c in df.columns if 'name' in c.lower() or 'city' in c.lower()), None)
                                if name_col:
                                    text_content += f"Highest {col}: {top_record[name_col]} ({series.max()}). "
                                    text_content += f"Lowest {col}: {bottom_record[name_col]} ({series.min()}). "
                            text_content += "\\n"
                    
                    # Categorical summaries
                    categorical_columns = df.select_dtypes(include=['object']).columns
                    for col in categorical_columns:
                        if col.lower() in ['state', 'district', 'region', 'category']:
                            series = df[col].dropna()
                            if not series.empty:
                                value_counts = series.value_counts()
                                text_content += f"BREAKDOWN: {sheet_name} has {len(value_counts)} different {col.lower()}s. "
                                top_3 = value_counts.head(3)
                                text_content += f"Most common: {', '.join(f'{val} ({count} records)' for val, count in top_3.items())}\\n"
                    
                    text_content += "\\n" + "="*60 + "\\n"
                    
                except Exception as e:
                    self.logger.warning(f"Error processing sheet '{sheet_name}': {e}")
                    continue
            
            if not text_content.strip():
                return f"Excel file: {os.path.basename(file_path)} (No data found in any sheets)"
            
            # Add comprehensive file summary
            summary = f"\\n=== EXCEL FILE: {os.path.basename(file_path)} ===\\n"
            summary += f"OVERVIEW: This Excel file contains {total_rows} records across {len(excel_file.sheet_names)} sheets. "
            summary += f"Each record represents a complete data entry with multiple attributes that can be searched and analyzed. "
            summary += f"The data is structured to answer questions about specific entities, comparisons, rankings, and statistical analysis.\\n"
            summary += f"SHEETS: {', '.join(excel_file.sheet_names)}\\n"
            summary += f"QUERYABLE: You can ask about specific records, comparisons between entities, statistical analysis, and relationships in the data.\\n\\n"
            
            final_content = summary + text_content
            
            self.logger.info(f"Successfully extracted {len(final_content)} characters from Excel with enhanced row-based preprocessing ({total_rows} records)")
            return final_content.strip()
            
        except Exception as e:
            self.logger.error(f"Error reading Excel {file_path}: {e}")
            raise Exception(f"Error reading Excel: {e}")
    
    def _process_image_file(self, file_path: str) -> str:
        """
        Enhanced image OCR processing with better error handling.
        
        Args:
            file_path (str): Path to the image file
        
        Returns:
            str: Extracted text content
        """
        try:
            self.logger.info(f"Processing Image: {os.path.basename(file_path)}")
            
            # Open and process the image
            image = Image.open(file_path)
            
            # Get image info
            image_info = f"Image: {os.path.basename(file_path)}\\n"
            image_info += f"Size: {image.size[0]}x{image.size[1]} pixels\\n"
            image_info += f"Mode: {image.mode}\\n"
            image_info += f"Format: {image.format}\\n"
            
            # Check if Tesseract is available
            if not TESSERACT_AVAILABLE:
                self.logger.warning(f"Tesseract OCR not available for {file_path}")
                return image_info + "\\nOCR Status: Tesseract not available - please install Tesseract to extract text from images\\n"
            
            # Preprocess image for better OCR
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Use Tesseract OCR to extract text
            self.logger.info("Running OCR on image...")
            ocr_config = '--oem 3 --psm 6'  # Use LSTM OCR Engine Mode with uniform text blocks
            text_content = pytesseract.image_to_string(image, config=ocr_config)
            
            if not text_content or not text_content.strip():
                return image_info + "\\nOCR Status: No text detected in image\\n"
            
            # Format the result
            result = image_info + "\\n=== OCR Extracted Text ===\\n" + text_content.strip()
            
            self.logger.info(f"Successfully extracted {len(text_content.strip())} characters from image using OCR")
            return result
            
        except Exception as e:
            # If OCR fails, return basic image info
            self.logger.warning(f"OCR failed for {file_path}: {e}")
            try:
                image = Image.open(file_path)
                return f"Image file: {os.path.basename(file_path)} (Size: {image.size}, OCR extraction failed: {str(e)})"
            except:
                return f"Image file: {os.path.basename(file_path)} (Could not process image: {str(e)})"
    
    def get_file_statistics(self, results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Generate statistics from processed files.
        
        Args:
            results (List[Dict[str, Any]]): List of processed file results
        
        Returns:
            Dict[str, Any]: Statistics summary
        """
        total_files = len(results)
        successful_files = sum(1 for r in results if r['status'] == 'success')
        failed_files = total_files - successful_files
        
        # Count by file type
        file_types = {}
        total_content_length = 0
        
        for result in results:
            file_type = result['file_type']
            file_types[file_type] = file_types.get(file_type, 0) + 1
            
            if result['status'] == 'success':
                total_content_length += len(result['content'])
        
        return {
            'total_files': total_files,
            'successful_files': successful_files,
            'failed_files': failed_files,
            'file_types': file_types,
            'total_content_length': total_content_length,
            'average_content_length': total_content_length / successful_files if successful_files > 0 else 0
        }
